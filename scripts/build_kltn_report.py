from copy import deepcopy
from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Mau bao cao KLTN_IS.docx"
ALT_OUT = ROOT / "Mau bao cao KLTN_IS_hoan_chinh.docx"
ALT_OUT_2 = ROOT / "Mau bao cao KLTN_IS_hoan_chinh_final.docx"
ACADEMIC_OUT = ROOT / "Mau bao cao KLTN_IS_hoan_chinh_hoc_thuat.docx"
BACKUP = ROOT / "Mau bao cao KLTN_IS_backup_before_full_report.docx"
ASSET_DIR = ROOT / "report_assets"
DIAGRAM_DIR = ASSET_DIR / "diagrams"
SCREENSHOT_DIR = ASSET_DIR / "screenshots"


def set_run_font(run, size=13, bold=False, italic=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_font(paragraph, size=13, bold=False, italic=False, color=None):
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, italic=italic, color=color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, size=12):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if bold else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def clone_paragraph_after(paragraph, text=None, style=None):
    new_p = deepcopy(paragraph._p)
    paragraph._p.addnext(new_p)
    p = paragraph._parent.paragraphs[-1]
    if text is not None:
        p.text = text
    if style:
        p.style = style
    return p


def add_field(paragraph, field_code):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)
    set_run_font(run, size=11)


def configure_document(doc):
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(3.5)
    section.right_margin = Cm(2)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Heading 1", 14, "000000"),
        ("Heading 2", 13, "000000"),
        ("Heading 3", 13, "000000"),
    ]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.space_before = Pt(12 if name == "Heading 1" else 6)
        style.paragraph_format.space_after = Pt(8)


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "Khóa luận tốt nghiệp - Hệ thống xét duyệt thi đua khen thưởng"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_font(header, size=10, italic=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Lê Trần Hoài Lộc - Trang ")
    add_field(footer, "PAGE")
    set_paragraph_font(footer, size=10)


def add_centered(paragraph, text, size=13, bold=False, spacing_after=6):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(spacing_after)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return paragraph


def add_body(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run_font(run, size=13)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        set_run_font(run, size=13)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        set_run_font(run, size=13)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.text = text
    set_paragraph_font(p, size=14 if level == 1 else 13, bold=True)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_table(doc, headers, rows, widths=None, title=None):
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        set_run_font(run, size=13, bold=True)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, size=12)
        set_cell_shading(hdr[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], str(text), size=12)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_run_font(run, size=12, italic=True)


def add_image(doc, path, caption, width_cm=15.5):
    path = Path(path)
    if not path.exists():
        return False
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)
    return True


def load_diagram_font(size=24, bold=False):
    candidates = [
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/timesbd.ttf" if bold else "C:/Windows/Fonts/times.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def text_size(draw, text, font):
    box = draw.textbbox((0, 0), text, font=font)
    return box[2] - box[0], box[3] - box[1]


def draw_centered(draw, box, text, font, fill="#111827"):
    x1, y1, x2, y2 = box
    w, h = text_size(draw, text, font)
    draw.text((x1 + (x2 - x1 - w) / 2, y1 + (y2 - y1 - h) / 2), text, font=font, fill=fill)


def draw_box(draw, box, text, font, fill="#F8FAFC", outline="#2563EB", radius=18):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=3)
    draw_centered(draw, box, text, font)


def draw_arrow(draw, start, end, fill="#334155", width=3):
    draw.line([start, end], fill=fill, width=width)
    x1, y1 = start
    x2, y2 = end
    dx, dy = x2 - x1, y2 - y1
    if abs(dx) >= abs(dy):
        direction = 1 if dx > 0 else -1
        head = [(x2, y2), (x2 - direction * 14, y2 - 8), (x2 - direction * 14, y2 + 8)]
    else:
        direction = 1 if dy > 0 else -1
        head = [(x2, y2), (x2 - 8, y2 - direction * 14), (x2 + 8, y2 - direction * 14)]
    draw.polygon(head, fill=fill)


def save_diagram(name, title, draw_fn):
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    path = DIAGRAM_DIR / name
    img = Image.new("RGB", (1600, 1000), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_diagram_font(34, bold=True)
    body_font = load_diagram_font(24)
    small_font = load_diagram_font(20)
    draw.text((50, 35), title, font=title_font, fill="#0F172A")
    draw_fn(draw, body_font, small_font)
    img.save(path)
    return path


def generate_diagrams():
    def use_case(draw, font, small):
        actors = [("Sinh viên/\nGiảng viên", 90, 210), ("Cán bộ", 90, 520), ("Hội đồng", 1320, 240), ("Admin", 1320, 560)]
        for label, x, y in actors:
            draw.ellipse((x + 40, y, x + 90, y + 50), outline="#0F172A", width=3)
            draw.line((x + 65, y + 50, x + 65, y + 120), fill="#0F172A", width=3)
            draw.line((x + 30, y + 78, x + 100, y + 78), fill="#0F172A", width=3)
            draw.line((x + 65, y + 120, x + 30, y + 170), fill="#0F172A", width=3)
            draw.line((x + 65, y + 120, x + 100, y + 170), fill="#0F172A", width=3)
            draw.multiline_text((x, y + 185), label, font=small, fill="#0F172A", align="center")
        use_cases = [
            ("Đăng nhập", 430, 150), ("Tạo hồ sơ", 430, 280), ("Upload minh chứng", 430, 410),
            ("Nộp hồ sơ", 430, 540), ("Duyệt đơn vị/khoa", 760, 300), ("Chấm điểm cấp trường", 760, 450),
            ("Quản lý tiêu chí/ý nhỏ", 1060, 260), ("Audit & dashboard", 1060, 430),
            ("Quét mã độc", 760, 610), ("Gửi thông báo", 1060, 610),
        ]
        for label, x, y in use_cases:
            draw.ellipse((x, y, x + 270, y + 82), fill="#EFF6FF", outline="#2563EB", width=3)
            draw_centered(draw, (x, y, x + 270, y + 82), label, small)
        links = [
            ((190, 300), (430, 190)), ((190, 330), (430, 320)), ((190, 360), (430, 450)), ((190, 390), (430, 580)),
            ((1320, 390), (1030, 490)), ((1320, 720), (1330, 305)), ((1320, 720), (1330, 475)),
            ((190, 610), (760, 340)), ((1320, 430), (1030, 490)),
        ]
        for start, end in links:
            draw.line([start, end], fill="#64748B", width=2)
        draw.line((565, 492, 870, 650), fill="#94A3B8", width=2)
        draw.text((620, 555), "<<include>>", font=small, fill="#475569")
        draw.line((895, 492, 1150, 650), fill="#94A3B8", width=2)
        draw.text((940, 555), "<<include>>", font=small, fill="#475569")

    def activity(draw, font, small):
        steps = [
            ("Bắt đầu", 650, 120), ("Tạo hồ sơ", 650, 230), ("Upload minh chứng", 650, 340),
            ("ClamAV quét file", 650, 450), ("Nộp hồ sơ", 650, 560),
            ("Duyệt đơn vị", 310, 690), ("Duyệt khoa", 650, 690), ("Hội đồng chấm điểm", 990, 690),
            ("Kết thúc", 650, 840),
        ]
        for label, x, y in steps:
            fill = "#DCFCE7" if label in ("Bắt đầu", "Kết thúc") else "#F8FAFC"
            draw_box(draw, (x, y, x + 260, y + 70), label, small, fill=fill, outline="#16A34A" if fill == "#DCFCE7" else "#2563EB")
        path = [(780, 190), (780, 230), (780, 300), (780, 340), (780, 410), (780, 450), (780, 520), (780, 560), (780, 630)]
        for a, b in zip(path, path[1:]):
            draw_arrow(draw, a, b)
        draw_arrow(draw, (760, 630), (440, 690))
        draw_arrow(draw, (780, 630), (780, 690))
        draw_arrow(draw, (800, 630), (1120, 690))
        draw_arrow(draw, (440, 760), (760, 840))
        draw_arrow(draw, (780, 760), (780, 840))
        draw_arrow(draw, (1120, 760), (900, 840))
        draw.text((930, 540), "Nếu file nhiễm mã độc: từ chối upload", font=small, fill="#B91C1C")

    def erd(draw, font, small):
        boxes = {
            "User": (80, 170), "Nomination": (430, 170), "Evidence": (820, 150), "ReviewStep": (820, 360),
            "Criteria": (80, 520), "CriteriaSubItem": (430, 520), "NominationItem": (820, 570),
            "AuditLog": (1180, 250), "Notification": (1180, 500),
        }
        fields = {
            "User": "id, email, role\nfullName, department",
            "Nomination": "id, ownerId, status\ntotalScore, submittedAt",
            "Evidence": "id, nominationId\nscanStatus, sha256",
            "ReviewStep": "id, nominationId\nreviewerId, decision",
            "Criteria": "id, name, maxScore\nlevel, targetRole",
            "CriteriaSubItem": "id, criteriaId\nname, maxScore",
            "NominationItem": "id, nominationId\ncriteriaId, score",
            "AuditLog": "id, actorId\naction, createdAt",
            "Notification": "id, userId\nmessage, readAt",
        }
        for name, (x, y) in boxes.items():
            draw.rounded_rectangle((x, y, x + 290, y + 135), radius=14, fill="#F8FAFC", outline="#2563EB", width=3)
            draw.rectangle((x, y, x + 290, y + 42), fill="#DBEAFE", outline="#2563EB", width=0)
            draw_centered(draw, (x, y, x + 290, y + 42), name, small)
            draw.multiline_text((x + 18, y + 58), fields[name], font=small, fill="#334155", spacing=5)
        rels = [
            ((370, 235), (430, 235), "1..n"), ((720, 235), (820, 210), "1..n"),
            ((720, 255), (820, 425), "1..n"), ((370, 585), (430, 585), "1..n"),
            ((720, 610), (820, 635), "1..n"), ((370, 285), (1180, 320), "1..n"),
            ((370, 265), (1180, 565), "1..n"), ((720, 235), (820, 635), "1..n"),
        ]
        for start, end, label in rels:
            draw_arrow(draw, start, end, width=2)
            mx, my = (start[0] + end[0]) / 2, (start[1] + end[1]) / 2
            draw.text((mx, my), label, font=small, fill="#475569")

    def sequence(draw, font, small):
        actors = ["Người dùng", "Frontend", "Backend API", "ClamAV", "Database", "Email"]
        xs = [120, 360, 620, 880, 1120, 1370]
        for actor, x in zip(actors, xs):
            draw_box(draw, (x - 80, 130, x + 80, 190), actor, small, fill="#EFF6FF")
            draw.line((x, 190, x, 900), fill="#CBD5E1", width=2)
        messages = [
            (120, 360, 250, "Chọn file / nộp hồ sơ"),
            (360, 620, 340, "POST upload/submit"),
            (620, 880, 430, "Gửi stream file"),
            (880, 620, 520, "CLEAN / INFECTED"),
            (620, 1120, 610, "Lưu hồ sơ, điểm, audit"),
            (620, 1370, 700, "Gửi thông báo"),
            (620, 360, 790, "Trả kết quả"),
        ]
        for x1, x2, y, label in messages:
            draw_arrow(draw, (x1, y), (x2, y), width=3)
            draw.text((min(x1, x2) + 15, y - 32), label, font=small, fill="#0F172A")

    def deployment(draw, font, small):
        boxes = [
            ("Trình duyệt người dùng\nReactJS/Vite", 120, 240, "#EFF6FF"),
            ("Backend API\nNode.js/Express", 520, 240, "#F0FDF4"),
            ("PostgreSQL\nDữ liệu nghiệp vụ", 980, 160, "#FFF7ED"),
            ("ClamAV daemon\nTCP 3310", 980, 360, "#FEF2F2"),
            ("SMTP service\nEmail thông báo", 980, 560, "#F8FAFC"),
            ("Upload storage\nquarantine/evidence", 520, 560, "#F8FAFC"),
        ]
        for label, x, y, fill in boxes:
            draw_box(draw, (x, y, x + 300, y + 115), label, small, fill=fill)
        draw_arrow(draw, (420, 300), (520, 300))
        draw.text((430, 260), "HTTPS/API", font=small, fill="#475569")
        draw_arrow(draw, (820, 285), (980, 210))
        draw_arrow(draw, (820, 320), (980, 420))
        draw_arrow(draw, (670, 355), (670, 560))
        draw_arrow(draw, (820, 335), (980, 620))

    paths = {
        "use_case": save_diagram("use_case_diagram.png", "Use Case Diagram - Hệ thống xét duyệt thi đua", use_case),
        "activity": save_diagram("activity_workflow.png", "Activity Diagram - Quy trình tạo và duyệt hồ sơ", activity),
        "erd": save_diagram("erd_database.png", "ERD - Mô hình dữ liệu chính", erd),
        "sequence": save_diagram("sequence_upload_approval.png", "Sequence Diagram - Upload, nộp và duyệt hồ sơ", sequence),
        "deployment": save_diagram("deployment_architecture.png", "Deployment Diagram - Kiến trúc triển khai", deployment),
    }
    return paths


def add_cover(doc):
    add_centered(doc.add_paragraph(), "BỘ CÔNG THƯƠNG", 13, True, 0)
    add_centered(doc.add_paragraph(), "TRƯỜNG ĐẠI HỌC CÔNG NGHIỆP TP. HỒ CHÍ MINH", 13, True, 0)
    add_centered(doc.add_paragraph(), "KHOA CÔNG NGHỆ THÔNG TIN", 13, True, 24)
    for _ in range(4):
        doc.add_paragraph()
    add_centered(doc.add_paragraph(), "LÊ TRẦN HOÀI LỘC", 14, True, 18)
    add_centered(doc.add_paragraph(), "XÂY DỰNG HỆ THỐNG XÉT DUYỆT", 16, True, 0)
    add_centered(doc.add_paragraph(), "THI ĐUA KHEN THƯỞNG CẤP TRƯỜNG", 16, True, 18)
    add_centered(doc.add_paragraph(), "BÁO CÁO KHÓA LUẬN TỐT NGHIỆP NĂM 2026", 14, True, 18)
    add_centered(doc.add_paragraph(), "Ngành: Hệ Thống Thông Tin", 13, False, 12)
    add_centered(doc.add_paragraph(), "Giảng viên hướng dẫn: TS. Võ Ngọc Tấn Phước", 13, False, 24)
    for _ in range(6):
        doc.add_paragraph()
    add_centered(doc.add_paragraph(), "THÀNH PHỐ HỒ CHÍ MINH, THÁNG 05 NĂM 2026", 13, True, 0)
    doc.add_page_break()


def add_front_matter(doc):
    add_heading(doc, "ABSTRACT", 1)
    for text in [
        "Topic: Design and Implementation of a University Emulation Award Review System.",
        "This thesis addresses limitations in manual and fragmented emulation-award review processes at university level, including inconsistent scoring, delayed approvals, difficult evidence tracking, and low traceability.",
        "The project proposes and implements a web-based information system with role-based workflows for students, lecturers, unit officers, school-level council members, and administrators.",
        "The solution includes nomination lifecycle management, multi-level approval, criteria configuration, detailed scoring by sub-items, evidence upload, malware scanning via ClamAV, notifications, operational dashboards, overdue reminders, reassignment, and audit logging.",
        "The system is built with ReactJS for the frontend, Node.js and Express for backend services, and PostgreSQL with Prisma ORM for data management. Security controls include JWT authentication, role-based authorization, file signature validation, antivirus scanning, controlled evidence download, and audit trails.",
        "Testing results show stable operation of the core workflows: nomination submission, staged review, reassignment, secure evidence handling, malware detection, scoring, and dashboard reporting.",
        "Keywords: emulation award system, workflow approval, role-based access control, secure file upload, ClamAV, audit log, ReactJS, Node.js, PostgreSQL.",
    ]:
        add_body(doc, text)
    doc.add_page_break()

    add_heading(doc, "TÓM TẮT", 1)
    for text in [
        "Tên đề tài: Xây dựng hệ thống xét duyệt thi đua khen thưởng cấp trường.",
        "Đề tài được thực hiện nhằm giải quyết các bất cập của quy trình xét duyệt thi đua thủ công như chậm xử lý, thiếu minh bạch, khó theo dõi tiến độ, khó kiểm soát minh chứng và thiếu đồng nhất khi chấm điểm.",
        "Khóa luận đề xuất và hiện thực một hệ thống web quản lý xuyên suốt vòng đời hồ sơ thi đua, hỗ trợ nhiều vai trò gồm sinh viên, giảng viên, cán bộ, hội đồng và quản trị viên.",
        "Hệ thống cung cấp các chức năng chính: tạo hồ sơ, nộp minh chứng, duyệt nhiều cấp, cấu hình tiêu chí lớn và ý nhỏ, chấm điểm chi tiết, phân công lại người duyệt, nhắc việc quá hạn, thống kê vận hành và lưu vết audit.",
        "Giải pháp kỹ thuật sử dụng ReactJS cho giao diện, Node.js/Express cho dịch vụ backend, PostgreSQL và Prisma cho dữ liệu. Các cơ chế an toàn bao gồm xác thực JWT, phân quyền theo vai trò, kiểm tra chữ ký tệp, quét mã độc ClamAV và kiểm soát tải xuống minh chứng.",
        "Kết quả kiểm thử cho thấy hệ thống vận hành ổn định theo đúng quy trình nghiệp vụ, đáp ứng tốt nhu cầu số hóa công tác thi đua khen thưởng trong môi trường đại học.",
        "Từ khóa: hệ thống thi đua khen thưởng, quy trình duyệt nhiều cấp, phân quyền vai trò, chấm điểm theo ý nhỏ, quét mã độc ClamAV, nhật ký audit.",
    ]:
        add_body(doc, text)
    doc.add_page_break()

    add_heading(doc, "LỜI CẢM ƠN", 1)
    for text in [
        "Em xin chân thành cảm ơn Ban Giám hiệu Trường Đại học Công nghiệp TP. Hồ Chí Minh và quý thầy cô Khoa Công nghệ Thông tin đã tạo điều kiện học tập, nghiên cứu và thực hiện khóa luận tốt nghiệp.",
        "Em xin bày tỏ lòng biết ơn sâu sắc đến TS. Võ Ngọc Tấn Phước đã tận tình hướng dẫn, góp ý chuyên môn và định hướng xuyên suốt quá trình thực hiện đề tài.",
        "Em cũng xin cảm ơn các anh chị, bạn bè và những người đã hỗ trợ góp ý nghiệp vụ, kiểm thử hệ thống và chia sẻ kinh nghiệm thực tế để hoàn thiện sản phẩm.",
        "Mặc dù đã nỗ lực trong quá trình thực hiện, báo cáo khó tránh khỏi thiếu sót. Em rất mong nhận được ý kiến đóng góp của quý thầy cô để tiếp tục hoàn thiện đề tài trong các giai đoạn tiếp theo.",
    ]:
        add_body(doc, text)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Người thực hiện đề tài\nLÊ TRẦN HOÀI LỘC")
    set_run_font(run, size=13, bold=True)
    doc.add_page_break()


def add_toc(doc):
    add_heading(doc, "MỤC LỤC", 1)
    toc = [
        "CHƯƠNG 1. GIỚI THIỆU ĐỀ TÀI",
        "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ",
        "CHƯƠNG 3. PHÂN TÍCH YÊU CẦU HỆ THỐNG",
        "CHƯƠNG 4. THIẾT KẾ HỆ THỐNG",
        "CHƯƠNG 5. GIẢI PHÁP CÔNG NGHỆ VÀ AN TOÀN HỆ THỐNG",
        "CHƯƠNG 6. HIỆN THỰC VÀ TRIỂN KHAI",
        "CHƯƠNG 7. KIỂM THỬ VÀ ĐÁNH GIÁ",
        "CHƯƠNG 8. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN",
        "TÀI LIỆU THAM KHẢO",
        "PHỤ LỤC",
    ]
    for item in toc:
        p = doc.add_paragraph()
        run = p.add_run(item)
        set_run_font(run, size=13)
    doc.add_page_break()


def add_chapter_1(doc):
    add_heading(doc, "CHƯƠNG 1. GIỚI THIỆU ĐỀ TÀI", 1)
    add_heading(doc, "1.1. Bối cảnh và lý do chọn đề tài", 2)
    for text in [
        "Trong môi trường đại học, công tác thi đua khen thưởng giữ vai trò quan trọng trong việc ghi nhận thành tích học tập, nghiên cứu khoa học, hoạt động phong trào và đóng góp của sinh viên, giảng viên cũng như các đơn vị trực thuộc. Tuy nhiên, tại nhiều đơn vị, quy trình tiếp nhận và xét duyệt hồ sơ vẫn phụ thuộc nhiều vào biểu mẫu rời rạc, trao đổi thủ công và xử lý qua nhiều kênh khác nhau.",
        "Cách làm truyền thống khiến quá trình tổng hợp minh chứng, kiểm tra điều kiện, chấm điểm và phê duyệt qua nhiều cấp dễ phát sinh sai sót. Người nộp hồ sơ khó theo dõi trạng thái, cán bộ duyệt khó kiểm soát hồ sơ tồn đọng, còn cấp quản lý thiếu số liệu tức thời để đánh giá hiệu quả vận hành.",
        "Từ thực tế đó, đề tài xây dựng hệ thống xét duyệt thi đua khen thưởng cấp trường nhằm số hóa quy trình nghiệp vụ, chuẩn hóa tiêu chí đánh giá, tăng tính minh bạch và hỗ trợ quản lý tập trung toàn bộ vòng đời hồ sơ.",
    ]:
        add_body(doc, text)
    add_heading(doc, "1.2. Mục tiêu đề tài", 2)
    add_bullets(doc, [
        "Xây dựng hệ thống web hỗ trợ tạo, nộp, duyệt và chấm điểm hồ sơ thi đua khen thưởng.",
        "Thiết kế quy trình duyệt nhiều cấp gồm cấp đơn vị, cấp khoa và cấp trường.",
        "Cho phép cấu hình tiêu chí lớn, ý nhỏ và chấm điểm chi tiết theo từng mục.",
        "Bổ sung cơ chế bảo mật file minh chứng bằng kiểm tra định dạng, chữ ký tệp và quét mã độc ClamAV.",
        "Xây dựng dashboard vận hành, nhắc việc quá hạn, phân công lại người duyệt và nhật ký audit.",
    ])
    add_heading(doc, "1.3. Phạm vi nghiên cứu", 2)
    add_body(doc, "Đề tài tập trung vào quy trình xét duyệt thi đua trong phạm vi trường đại học, với các đối tượng chính là sinh viên, giảng viên, cán bộ phụ trách, hội đồng thi đua và quản trị viên. Hệ thống không đi sâu vào tích hợp với phần mềm quản lý đào tạo hiện hữu, nhưng được thiết kế theo hướng có thể mở rộng kết nối trong tương lai.")
    add_heading(doc, "1.4. Phương pháp thực hiện", 2)
    add_numbered(doc, [
        "Khảo sát quy trình nghiệp vụ và xác định các actor tham gia hệ thống.",
        "Phân tích yêu cầu chức năng, phi chức năng và ràng buộc bảo mật.",
        "Thiết kế cơ sở dữ liệu, API, phân quyền và giao diện người dùng.",
        "Hiện thực hệ thống bằng ReactJS, Node.js/Express, PostgreSQL và Prisma.",
        "Kiểm thử theo luồng nghiệp vụ thực tế và kiểm thử bảo mật upload file.",
    ])


def add_chapter_2(doc):
    add_heading(doc, "CHƯƠNG 2. CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ", 1)
    add_heading(doc, "2.1. Tổng quan hệ thống thông tin quản lý", 2)
    add_body(doc, "Hệ thống thông tin quản lý là tập hợp các thành phần phần mềm, dữ liệu, quy trình và con người nhằm hỗ trợ việc thu thập, xử lý, lưu trữ và cung cấp thông tin phục vụ ra quyết định. Trong đề tài này, hệ thống thông tin quản lý được áp dụng để chuẩn hóa quy trình xét duyệt thi đua, giảm xử lý thủ công và tăng khả năng giám sát.")
    add_heading(doc, "2.2. Quy trình duyệt nhiều cấp", 2)
    add_body(doc, "Quy trình duyệt nhiều cấp là mô hình trong đó một hồ sơ phải đi qua các bước phê duyệt tuần tự. Mỗi cấp có trách nhiệm kiểm tra một nhóm điều kiện nhất định trước khi chuyển hồ sơ lên cấp tiếp theo. Mô hình này phù hợp với nghiệp vụ thi đua khen thưởng vì hồ sơ cần được kiểm tra từ đơn vị cơ sở đến cấp hội đồng.")
    add_heading(doc, "2.3. Phân quyền theo vai trò", 2)
    add_body(doc, "Role-Based Access Control (RBAC) là cơ chế phân quyền dựa trên vai trò người dùng. Mỗi vai trò được gán một tập quyền cụ thể, giúp đảm bảo người dùng chỉ được thực hiện các thao tác phù hợp với nhiệm vụ của mình. Hệ thống áp dụng RBAC cho các vai trò ADMIN, CANBO, GIANGVIEN, SINHVIEN và HOIDONG.")
    add_table(doc, ["Vai trò", "Mô tả trách nhiệm"], [
        ["ADMIN", "Quản trị người dùng, cấu hình tiêu chí, theo dõi dashboard, reassign reviewer, kích hoạt quét lại file."],
        ["SINHVIEN", "Tạo hồ sơ, upload minh chứng, nộp hồ sơ, theo dõi trạng thái và mở lại hồ sơ bị từ chối."],
        ["GIANGVIEN", "Tạo hồ sơ thi đua dành cho giảng viên và theo dõi kết quả xét duyệt."],
        ["CANBO", "Duyệt hồ sơ ở cấp đơn vị hoặc cấp khoa, nhập nhận xét, từ chối hoặc chuyển cấp."],
        ["HOIDONG", "Duyệt cấp trường, chấm điểm cuối cùng và công nhận hồ sơ."],
    ], title="Bảng 2.1. Vai trò người dùng trong hệ thống")
    add_heading(doc, "2.4. Công nghệ sử dụng", 2)
    add_table(doc, ["Thành phần", "Công nghệ", "Vai trò"], [
        ["Frontend", "ReactJS, Vite", "Xây dựng giao diện và tương tác người dùng."],
        ["Backend", "Node.js, Express", "Cung cấp API, xử lý nghiệp vụ và phân quyền."],
        ["Database", "PostgreSQL", "Lưu trữ người dùng, hồ sơ, tiêu chí, điểm, audit và minh chứng."],
        ["ORM", "Prisma", "Mô hình hóa dữ liệu, migration và truy vấn database."],
        ["Bảo mật file", "ClamAV", "Quét mã độc đối với file minh chứng được tải lên."],
        ["Email", "Nodemailer", "Gửi thông báo duyệt, từ chối và đặt lại mật khẩu."],
    ], title="Bảng 2.2. Công nghệ chính của hệ thống")
    add_heading(doc, "2.5. Nghiên cứu liên quan", 2)
    add_body(doc, "Các hệ thống xét duyệt điện tử trong môi trường giáo dục thường được xây dựng xoay quanh ba nhóm vấn đề: quản lý quy trình phê duyệt nhiều cấp, kiểm soát quyền truy cập theo vai trò và bảo đảm tính toàn vẹn của hồ sơ minh chứng. Những nhóm vấn đề này tương ứng với các hướng nghiên cứu và thực hành phổ biến gồm workflow approval system, role-based access control và secure file upload.")
    add_table(doc, ["Nhóm giải pháp", "Đặc điểm", "Ý nghĩa đối với đề tài"], [
        ["Workflow approval system", "Mô hình hóa quy trình thành các bước có trạng thái, người phụ trách, quyết định và hạn xử lý.", "Làm cơ sở thiết kế ReviewStep, nhắc việc quá hạn và chuyển cấp duyệt."],
        ["RBAC", "Gán quyền theo vai trò thay vì gán trực tiếp từng quyền cho từng cá nhân.", "Giúp phân tách rõ Sinh viên/Giảng viên, Cán bộ, Hội đồng và Admin."],
        ["Secure upload system", "Kiểm tra định dạng, chữ ký file, đổi tên file, lưu cách ly và quét mã độc.", "Giảm rủi ro khi tiếp nhận minh chứng do người dùng tải lên."],
        ["Audit trail", "Ghi lại actor, hành động, thời gian và đối tượng bị tác động.", "Tăng khả năng truy vết và minh bạch khi xử lý hồ sơ thi đua."],
    ], title="Bảng 2.3. Tổng hợp nghiên cứu và giải pháp liên quan")


def add_chapter_3(doc):
    add_heading(doc, "CHƯƠNG 3. PHÂN TÍCH YÊU CẦU HỆ THỐNG", 1)
    add_heading(doc, "3.1. Phân tích actor", 2)
    add_body(doc, "Hệ thống gồm năm nhóm actor chính. Mỗi actor có phạm vi thao tác riêng, được ràng buộc bởi cơ chế xác thực và phân quyền. Việc phân chia rõ vai trò giúp quy trình duyệt minh bạch và hạn chế thao tác sai phạm.")
    add_table(doc, ["Actor", "Use case chính"], [
        ["Sinh viên", "Đăng ký, đăng nhập, tạo hồ sơ, upload minh chứng, nộp duyệt, xem trạng thái."],
        ["Giảng viên", "Tạo hồ sơ thi đua, nhập thông tin minh chứng, nộp duyệt."],
        ["Cán bộ", "Xem hồ sơ được phân công, duyệt, từ chối, nhập nhận xét, nhận nhắc việc quá hạn."],
        ["Hội đồng", "Chấm điểm cấp trường, phê duyệt cuối cùng hoặc từ chối hồ sơ."],
        ["Admin", "Quản lý tài khoản, tiêu chí, ý nhỏ, dashboard, reassign, quét lại file, cấu hình hệ thống."],
    ], title="Bảng 3.1. Actor và use case chính")
    add_heading(doc, "3.2. Yêu cầu chức năng", 2)
    add_bullets(doc, [
        "Người dùng có thể đăng nhập, cập nhật thông tin cá nhân và đổi mật khẩu.",
        "Admin có thể tạo, sửa, xóa người dùng và phân quyền vai trò.",
        "Admin hoặc hội đồng có thể cấu hình tiêu chí lớn và ý nhỏ để chấm điểm.",
        "Sinh viên và giảng viên có thể tạo hồ sơ thi đua và tải lên minh chứng.",
        "Hệ thống tự tạo các bước duyệt khi hồ sơ được nộp.",
        "Cán bộ và hội đồng chỉ xử lý các phiên duyệt được phân công.",
        "Hội đồng có thể chấm điểm chi tiết và hệ thống tự cộng tổng điểm.",
        "Admin có thể phân công lại reviewer và kích hoạt quét lại file chờ quét.",
        "Hệ thống có dashboard thống kê hồ sơ, file chờ quét, file nhiễm mã độc và review quá hạn.",
    ])
    add_heading(doc, "3.3. Yêu cầu phi chức năng", 2)
    add_table(doc, ["Nhóm yêu cầu", "Mô tả"], [
        ["Bảo mật", "Xác thực JWT, phân quyền RBAC, kiểm soát tải file, quét mã độc."],
        ["Tin cậy", "Dữ liệu được lưu trong PostgreSQL, thao tác quan trọng có audit log."],
        ["Dễ dùng", "Giao diện phân vai rõ ràng, thao tác tạo hồ sơ và duyệt đơn giản."],
        ["Mở rộng", "Thiết kế API và schema cho phép thêm cấp duyệt, báo cáo và tích hợp SSO."],
        ["Hiệu năng", "Các truy vấn chính được giới hạn theo vai trò và trạng thái để giảm tải."],
    ], title="Bảng 3.2. Yêu cầu phi chức năng")
    add_heading(doc, "3.4. Quy trình nghiệp vụ", 2)
    add_numbered(doc, [
        "Người nộp tạo hồ sơ, chọn tiêu chí và tải minh chứng.",
        "Hệ thống kiểm tra định dạng file, chữ ký file và quét mã độc.",
        "Người nộp gửi hồ sơ sang trạng thái chờ duyệt.",
        "Cán bộ cấp đơn vị kiểm tra điều kiện và duyệt hoặc từ chối.",
        "Cán bộ cấp khoa tiếp tục kiểm tra và duyệt hoặc từ chối.",
        "Hội đồng cấp trường chấm điểm, duyệt cuối cùng và công nhận hồ sơ.",
        "Hệ thống gửi thông báo, email và lưu nhật ký thao tác.",
    ])
    add_heading(doc, "3.5. So sánh quy trình thủ công và hệ thống đề xuất", 2)
    add_body(doc, "Để làm rõ giá trị nghiên cứu và ứng dụng, đề tài đối chiếu quy trình thủ công hiện tại với hệ thống đề xuất theo các tiêu chí nghiệp vụ, quản trị và bảo mật. Kết quả cho thấy hệ thống đề xuất không chỉ thay thế thao tác giấy tờ mà còn bổ sung khả năng chuẩn hóa, đo lường và truy vết.")
    add_table(doc, ["Tiêu chí", "Quy trình thủ công", "Hệ thống đề xuất"], [
        ["Theo dõi tiến độ", "Phụ thuộc email, file rời hoặc trao đổi trực tiếp.", "Mỗi hồ sơ có trạng thái, cấp duyệt và người xử lý rõ ràng."],
        ["Minh bạch chấm điểm", "Điểm tổng thường khó giải thích chi tiết.", "Tiêu chí lớn được chia thành ý nhỏ, có điểm từng mục và tổng điểm tự động."],
        ["Thời gian xử lý", "Dễ chậm do chuyển hồ sơ thủ công.", "Tự động chuyển bước, thông báo và nhắc việc quá hạn."],
        ["Bảo mật minh chứng", "File có thể được chia sẻ rời rạc, khó kiểm soát.", "File được kiểm tra định dạng, quét ClamAV và chỉ tải qua API có phân quyền."],
        ["Audit", "Khó xác định ai thao tác và thao tác lúc nào.", "AuditLog ghi nhận hành động quan trọng để truy vết."],
        ["Báo cáo", "Tổng hợp thủ công, dễ sai lệch.", "Dashboard và báo cáo dựa trên dữ liệu tập trung."],
    ], title="Bảng 3.3. So sánh quy trình thủ công và hệ thống đề xuất")
    add_heading(doc, "3.6. Use Case Diagram", 2)
    add_body(doc, "Use Case Diagram mô tả quan hệ giữa các actor và chức năng chính. Các quan hệ include được dùng cho các thao tác bắt buộc như xác thực, kiểm tra quyền, quét mã độc và gửi thông báo; các nhánh từ chối hoặc mở lại hồ sơ có thể được xem là luồng mở rộng của quy trình duyệt.")
    diagrams = generate_diagrams()
    add_image(doc, diagrams["use_case"], "Hình 3.1. Use Case Diagram của hệ thống", width_cm=15.5)


def add_chapter_4(doc):
    add_heading(doc, "CHƯƠNG 4. THIẾT KẾ HỆ THỐNG", 1)
    diagrams = generate_diagrams()
    add_heading(doc, "4.1. Kiến trúc tổng thể", 2)
    add_body(doc, "Hệ thống được thiết kế theo mô hình client-server. Frontend ReactJS chịu trách nhiệm hiển thị giao diện và gửi yêu cầu đến backend. Backend Express xử lý xác thực, phân quyền, nghiệp vụ hồ sơ, duyệt, chấm điểm, upload minh chứng và gửi email. PostgreSQL là nơi lưu trữ dữ liệu nghiệp vụ, còn ClamAV được dùng như dịch vụ quét mã độc độc lập.")
    add_table(doc, ["Lớp", "Thành phần", "Chức năng"], [
        ["Client", "ReactJS", "Hiển thị giao diện, gọi API, quản lý trạng thái phiên đăng nhập."],
        ["API", "Express Router", "Cung cấp endpoint cho auth, users, criteria, nominations, reviews, system."],
        ["Service", "Utils và jobs", "Gửi email, audit, scan file, scheduler quét lại file."],
        ["Database", "PostgreSQL", "Lưu dữ liệu người dùng, hồ sơ, duyệt, điểm, audit, minh chứng."],
        ["Security", "ClamAV", "Quét mã độc file minh chứng trước khi cho phép sử dụng."],
    ], title="Bảng 4.1. Kiến trúc logic của hệ thống")
    add_image(doc, diagrams["deployment"], "Hình 4.1. Deployment Diagram mô tả FE/BE/DB/ClamAV/SMTP", width_cm=15.5)
    add_heading(doc, "4.2. Thiết kế cơ sở dữ liệu", 2)
    add_body(doc, "Cơ sở dữ liệu được mô hình hóa bằng Prisma schema. Các bảng chính gồm User, Criteria, CriteriaSubItem, Nomination, NominationItem, ReviewStep, Evidence, ApprovalResult, Notification và AuditLog. Mối quan hệ giữa các bảng phản ánh vòng đời hồ sơ từ lúc tạo đến khi duyệt cuối cùng.")
    add_table(doc, ["Bảng", "Vai trò dữ liệu"], [
        ["User", "Lưu tài khoản, vai trò, thông tin cá nhân và đơn vị."],
        ["Criteria", "Lưu tiêu chí lớn, điểm tối đa, đối tượng áp dụng và cấp xét duyệt."],
        ["CriteriaSubItem", "Lưu các ý nhỏ của tiêu chí lớn để chấm điểm chi tiết."],
        ["Nomination", "Lưu hồ sơ thi đua, trạng thái, tổng điểm và người nộp."],
        ["NominationItem", "Lưu từng tiêu chí được chọn trong hồ sơ và điểm tương ứng."],
        ["ReviewStep", "Lưu các phiên duyệt theo cấp, người duyệt, hạn duyệt và quyết định."],
        ["Evidence", "Lưu thông tin file minh chứng, hash, trạng thái quét mã độc và mô tả."],
        ["AuditLog", "Lưu nhật ký thao tác quan trọng để phục vụ truy vết."],
    ], title="Bảng 4.2. Các bảng dữ liệu chính")
    add_image(doc, diagrams["erd"], "Hình 4.2. ERD và quan hệ cardinality chính", width_cm=15.5)
    add_heading(doc, "4.3. Activity Diagram quy trình nghiệp vụ", 2)
    add_body(doc, "Activity Diagram thể hiện luồng xử lý từ khi người dùng tạo hồ sơ, tải minh chứng, hệ thống quét mã độc, nộp hồ sơ và đi qua các cấp duyệt. Sơ đồ giúp chứng minh quy trình không chỉ được mô tả bằng chữ mà đã được mô hình hóa thành luồng nghiệp vụ có điểm bắt đầu, điều kiện rẽ nhánh và điểm kết thúc.")
    add_image(doc, diagrams["activity"], "Hình 4.3. Activity Diagram quy trình tạo và duyệt hồ sơ", width_cm=15.5)
    add_heading(doc, "4.4. Sequence Diagram", 2)
    add_body(doc, "Sequence Diagram mô tả tương tác theo thời gian giữa người dùng, frontend, backend, ClamAV, database và email service. Đây là cơ sở để kiểm tra tính đúng đắn của các API upload minh chứng, nộp hồ sơ, lưu dữ liệu duyệt và gửi thông báo.")
    add_image(doc, diagrams["sequence"], "Hình 4.4. Sequence Diagram upload, nộp và duyệt hồ sơ", width_cm=15.5)
    add_heading(doc, "4.5. Thiết kế API", 2)
    add_table(doc, ["Nhóm API", "Chức năng"], [
        ["/auth", "Đăng ký, đăng nhập, quên mật khẩu, đặt lại mật khẩu, thông tin cá nhân."],
        ["/users", "Quản lý người dùng dành cho admin."],
        ["/criteria", "Quản lý tiêu chí lớn và ý nhỏ."],
        ["/nominations", "Tạo, cập nhật, nộp hồ sơ, upload/tải/xóa minh chứng."],
        ["/reviews", "Xem hồ sơ chờ duyệt, quyết định duyệt, reassign, nhắc việc."],
        ["/system", "Dashboard vận hành, backup, quét lại file chờ quét."],
    ], title="Bảng 4.3. Nhóm API chính")
    add_heading(doc, "4.6. Thiết kế giao diện", 2)
    add_body(doc, "Giao diện được tổ chức theo nhóm chức năng phù hợp từng vai trò. Người nộp tập trung vào tạo và theo dõi hồ sơ; cán bộ và hội đồng tập trung vào danh sách hồ sơ chờ duyệt; admin có dashboard vận hành, quản lý người dùng và cấu hình tiêu chí. Cách tổ chức này giảm nhiễu thông tin và giúp người dùng thao tác đúng phạm vi quyền hạn.")


def add_chapter_5(doc):
    add_heading(doc, "CHƯƠNG 5. GIẢI PHÁP CÔNG NGHỆ VÀ AN TOÀN HỆ THỐNG", 1)
    add_heading(doc, "5.1. Xác thực và phân quyền", 2)
    add_body(doc, "Hệ thống sử dụng JWT để xác thực phiên đăng nhập. Sau khi đăng nhập thành công, token được gửi kèm trong các request cần bảo vệ. Backend kiểm tra token, tải thông tin người dùng từ database và áp dụng middleware phân quyền theo vai trò trước khi xử lý nghiệp vụ.")
    add_heading(doc, "5.2. Cơ chế upload file an toàn", 2)
    add_body(doc, "Minh chứng là thành phần quan trọng nhưng cũng có rủi ro bảo mật cao. Vì vậy, hệ thống áp dụng nhiều lớp bảo vệ: giới hạn đuôi file, kiểm tra MIME type, kiểm tra chữ ký file, đổi tên file bằng UUID, lưu tạm vào thư mục quarantine, quét ClamAV và chỉ cho phép tải xuống file đã đạt trạng thái CLEAN.")
    add_table(doc, ["Lớp kiểm soát", "Mục đích"], [
        ["Whitelist extension", "Chỉ cho phép PDF, DOCX, XLSX, PNG, JPG, JPEG và ZIP."],
        ["MIME type", "Giảm nguy cơ giả mạo định dạng ở mức trình duyệt/upload."],
        ["Magic bytes", "Xác thực nội dung file đúng với định dạng khai báo."],
        ["ClamAV", "Phát hiện malware, bao gồm mẫu EICAR trong kiểm thử."],
        ["Controlled download", "Không public thư mục upload, chỉ tải qua API có kiểm quyền."],
    ], title="Bảng 5.1. Các lớp bảo vệ file upload")
    add_heading(doc, "5.3. Trạng thái quét mã độc", 2)
    add_body(doc, "Mỗi file minh chứng có trạng thái scanStatus gồm PENDING_SCAN, CLEAN, INFECTED hoặc SCAN_ERROR. Khi ClamAV hoạt động ở chế độ strict, file không thể hoàn tất upload nếu không quét được. Scheduler nền hỗ trợ quét lại các file còn ở trạng thái chờ.")
    add_heading(doc, "5.4. Audit và giám sát", 2)
    add_body(doc, "Các thao tác quan trọng như tạo hồ sơ, duyệt, reassign, tải minh chứng, sửa tiêu chí và thay đổi ý nhỏ đều được ghi vào AuditLog. Dashboard vận hành giúp admin theo dõi số hồ sơ chờ duyệt, review quá hạn, file chờ quét, file nhiễm mã độc và lỗi quét.")


def add_chapter_6(doc):
    add_heading(doc, "CHƯƠNG 6. HIỆN THỰC VÀ TRIỂN KHAI", 1)
    add_heading(doc, "6.1. Hiện thực frontend", 2)
    add_body(doc, "Frontend được xây dựng bằng ReactJS và Vite. Các trang chính gồm LoginPage, DashboardPage, NominationsPage, CriteriaPage, ReviewsPage, UsersPage, ReportsPage và ProfilePage. AuthContext quản lý trạng thái đăng nhập và token, trong khi api/client.js cấu hình Axios để gọi backend.")
    add_heading(doc, "6.2. Hiện thực backend", 2)
    add_body(doc, "Backend được tổ chức theo mô hình route-service đơn giản. Mỗi nhóm nghiệp vụ có route riêng như auth.routes.js, users.routes.js, criteria.routes.js, nominations.routes.js, reviews.routes.js và system.routes.js. Middleware authenticate và authorize đảm bảo chỉ người dùng hợp lệ mới truy cập được chức năng tương ứng.")
    add_heading(doc, "6.3. Hiện thực nghiệp vụ chấm điểm", 2)
    add_body(doc, "Tiêu chí lớn có thể được chia thành nhiều ý nhỏ. Admin hoặc hội đồng cấu hình ý nhỏ với điểm tối đa, hệ thống kiểm soát tổng điểm ý nhỏ không vượt quá điểm tối đa của tiêu chí lớn. Khi hội đồng chấm, điểm từng ý nhỏ được cộng thành điểm tiêu chí và tiếp tục cộng vào tổng điểm hồ sơ.")
    add_heading(doc, "6.4. Triển khai ClamAV", 2)
    add_body(doc, "ClamAV được chạy dưới dạng clamd lắng nghe trên TCP port 3310. Backend kết nối đến clamd bằng giao thức INSTREAM để gửi nội dung file cần quét. Cấu hình môi trường gồm CLAMAV_HOST, CLAMAV_PORT và ALLOW_UNSCANNED_UPLOADS. Trong môi trường vận hành chính thức, hệ thống đặt ALLOW_UNSCANNED_UPLOADS=false để không chấp nhận file chưa được quét.")
    add_heading(doc, "6.5. Cấu hình và migration", 2)
    add_body(doc, "Prisma được sử dụng để quản lý schema và migration. Các migration quan trọng gồm bổ sung ý nhỏ tiêu chí, trạng thái quét minh chứng và hạn duyệt của ReviewStep. Sau mỗi thay đổi schema, hệ thống chạy prisma migrate và prisma generate để đồng bộ database với Prisma Client.")
    add_heading(doc, "6.6. Minh họa giao diện hệ thống", 2)
    add_body(doc, "Chương này cần có hình giao diện thật để chứng minh hệ thống đã được hiện thực. Các màn hình trọng tâm gồm đăng nhập, dashboard, tạo hồ sơ, upload minh chứng, quản lý tiêu chí, màn hình duyệt/chấm điểm, audit/dashboard vận hành và quản lý người dùng.")
    screenshots = [
        ("01_login.png", "Hình 6.1. Giao diện đăng nhập"),
        ("02_dashboard.png", "Hình 6.2. Dashboard tổng quan"),
        ("03_criteria_management.png", "Hình 6.3. Quản lý tiêu chí và ý nhỏ"),
        ("04_review_screen.png", "Hình 6.4. Màn hình duyệt và chấm điểm"),
        ("05_nomination_management.png", "Hình 6.5. Quản lý hồ sơ thi đua"),
        ("07_create_nomination_upload.png", "Hình 6.6. Tạo hồ sơ và upload minh chứng"),
    ]
    added = 0
    for file_name, caption in screenshots:
        if add_image(doc, SCREENSHOT_DIR / file_name, caption, width_cm=15.5):
            added += 1
    if added == 0:
        add_table(doc, ["Mã hình", "Màn hình cần chụp", "Mục đích minh họa"], [
            ["Hình 6.1", "Login", "Xác thực người dùng trước khi vào hệ thống."],
            ["Hình 6.2", "Dashboard", "Theo dõi số hồ sơ, review quá hạn, file chờ quét và cảnh báo vận hành."],
            ["Hình 6.3", "Create nomination", "Người dùng tạo hồ sơ thi đua và chọn tiêu chí."],
            ["Hình 6.4", "Upload evidence", "Minh họa whitelist định dạng và trạng thái quét mã độc."],
            ["Hình 6.5", "Criteria management", "Admin/Hội đồng quản lý tiêu chí lớn và ý nhỏ."],
            ["Hình 6.6", "Review screen", "Cán bộ/Hội đồng duyệt, nhận xét và chấm điểm."],
            ["Hình 6.7", "Audit log", "Truy vết thao tác quan trọng trong hệ thống."],
        ], title="Bảng 6.1. Danh mục screenshot cần đưa vào báo cáo khi chụp giao diện")


def add_chapter_7(doc):
    add_heading(doc, "CHƯƠNG 7. KIỂM THỬ VÀ ĐÁNH GIÁ", 1)
    add_heading(doc, "7.1. Kiểm thử chức năng theo actor", 2)
    add_body(doc, "Quá trình UAT được thực hiện theo từng actor chính. Kết quả kiểm thử cho thấy các luồng đăng nhập, tạo hồ sơ, upload minh chứng, nộp hồ sơ, duyệt cấp đơn vị, duyệt cấp khoa, reassign, duyệt cấp trường, chấm điểm và cập nhật trạng thái cuối đều hoạt động đúng.")
    add_table(doc, ["Ca kiểm thử", "Kết quả"], [
        ["Đăng nhập ADMIN, CANBO, HOIDONG, SINHVIEN", "Đạt"],
        ["Sinh viên upload minh chứng PDF sạch", "Đạt, scanStatus = CLEAN"],
        ["Sinh viên tạo và nộp hồ sơ", "Đạt"],
        ["Cán bộ cấp đơn vị duyệt", "Đạt"],
        ["Admin reassign reviewer cấp khoa", "Đạt"],
        ["Cán bộ được reassign duyệt cấp khoa", "Đạt"],
        ["Hội đồng chấm điểm và duyệt cấp trường", "Đạt"],
        ["Kiểm tra trạng thái cuối APPROVED", "Đạt"],
        ["Gửi nhắc việc quá hạn", "Đạt"],
        ["Dashboard vận hành", "Đạt"],
    ], title="Bảng 7.1. Kết quả UAT nghiệp vụ")
    add_heading(doc, "7.2. Kiểm thử bảo mật upload", 2)
    add_body(doc, "Hệ thống được kiểm thử bằng file sạch và mẫu EICAR. File sạch được ClamAV trả về trạng thái CLEAN, trong khi mẫu EICAR được phát hiện với kết quả INFECTED. Điều này chứng minh đường quét mã độc qua ClamAV hoạt động đúng ở mức kỹ thuật.")
    add_table(doc, ["Mẫu kiểm thử", "Kết quả mong đợi", "Kết quả thực tế"], [
        ["File văn bản/PDF sạch", "CLEAN", "CLEAN"],
        ["EICAR test file", "INFECTED", "INFECTED"],
        ["File sai chữ ký định dạng", "Bị từ chối", "Bị từ chối"],
    ], title="Bảng 7.2. Kiểm thử bảo mật file")
    add_heading(doc, "7.3. Testcase chi tiết", 2)
    add_body(doc, "Bảng testcase chi tiết giúp việc nghiệm thu có thể lặp lại, đo được và đối chiếu trực tiếp giữa input, expected output, actual output và trạng thái. Các testcase dưới đây tập trung vào luồng nghiệp vụ cốt lõi và các rủi ro cao của hệ thống.")
    add_table(doc, ["TC", "Chức năng", "Input", "Expected", "Actual", "Status"], [
        ["TC01", "Đăng nhập admin", "admin@iuh.edu.vn / 123456", "Vào dashboard admin", "Đăng nhập thành công", "Pass"],
        ["TC02", "Đăng nhập sinh viên", "hoailoc0505@gmail.com / 123456", "Vào màn hình hồ sơ cá nhân", "Đăng nhập thành công", "Pass"],
        ["TC03", "Tạo hồ sơ", "Tên hồ sơ, năm học, tiêu chí", "Hồ sơ ở trạng thái DRAFT", "Hồ sơ được tạo", "Pass"],
        ["TC04", "Upload PDF sạch", "File PDF hợp lệ", "scanStatus = CLEAN", "File được lưu và đánh dấu CLEAN", "Pass"],
        ["TC05", "Upload EICAR", "File chứa mẫu EICAR", "Từ chối upload/INFECTED", "ClamAV phát hiện INFECTED", "Pass"],
        ["TC06", "Upload sai định dạng", "File .exe hoặc sai magic bytes", "Từ chối file", "Backend trả lỗi định dạng", "Pass"],
        ["TC07", "Nộp hồ sơ", "Hồ sơ có minh chứng sạch", "Sinh review step cấp đơn vị", "Tạo bước duyệt", "Pass"],
        ["TC08", "Duyệt cấp đơn vị", "CANBO approve", "Chuyển sang cấp khoa", "Trạng thái cập nhật đúng", "Pass"],
        ["TC09", "Reassign reviewer", "Admin chọn reviewer mới", "ReviewStep đổi người xử lý", "Cập nhật thành công", "Pass"],
        ["TC10", "Hội đồng chấm điểm", "Điểm từng ý nhỏ", "Tổng điểm tự động", "Tổng điểm đúng theo ý nhỏ", "Pass"],
        ["TC11", "Từ chối hồ sơ", "Reviewer nhập lý do từ chối", "Hồ sơ REJECTED và có ghi chú", "Từ chối thành công", "Pass"],
        ["TC12", "Audit log", "Thao tác duyệt/reassign", "Có log actor, action, time", "Log được ghi", "Pass"],
    ], widths=[1.3, 2.7, 3.0, 3.2, 3.2, 1.3], title="Bảng 7.3. Testcase chi tiết")
    add_heading(doc, "7.4. Benchmark và đánh giá định lượng", 2)
    add_body(doc, "Đánh giá định lượng được thực hiện ở mức phù hợp với phạm vi khóa luận, tập trung vào các chỉ số có ảnh hưởng trực tiếp đến trải nghiệm và an toàn hệ thống. Một số số liệu là kết quả kiểm thử cục bộ, có thể thay đổi khi triển khai trên server thật.")
    add_table(doc, ["Chỉ số", "Kịch bản đo", "Kết quả tham chiếu", "Nhận xét"], [
        ["Thời gian scan file sạch", "Upload file minh chứng nhỏ dưới 2 MB", "Dưới 2 giây trong môi trường cục bộ", "Phù hợp với thao tác hồ sơ thông thường."],
        ["Phát hiện malware", "Upload mẫu EICAR", "Phát hiện 1/1 mẫu", "Xác nhận đường quét ClamAV hoạt động."],
        ["Số hồ sơ UAT", "Luồng tạo, nộp, duyệt nhiều cấp", "Ít nhất 1 hồ sơ đi đủ vòng đời", "Đủ chứng minh nghiệp vụ cốt lõi."],
        ["Response API đăng nhập", "POST /auth/login", "Phản hồi trong vài trăm ms ở môi trường local", "Phụ thuộc cấu hình máy và database."],
        ["Tải danh sách tiêu chí", "GET /criteria", "Hoạt động sau khi sửa lỗi hệ thống", "Đáp ứng nhu cầu chấm điểm."],
        ["Đồng thời cơ bản", "Nhiều actor đăng nhập theo UAT", "Không phát sinh lỗi phân quyền", "Cần stress test sâu hơn khi production."],
    ], title="Bảng 7.4. Benchmark và đánh giá định lượng")
    add_heading(doc, "7.5. Đánh giá kết quả", 2)
    add_body(doc, "Hệ thống đáp ứng được các mục tiêu chính của đề tài. Quy trình duyệt nhiều cấp giúp chuẩn hóa trách nhiệm từng actor. Cấu hình tiêu chí và ý nhỏ giúp việc chấm điểm minh bạch hơn. Dashboard và audit log giúp admin nắm được tình trạng vận hành. Cơ chế upload an toàn giảm rủi ro khi tiếp nhận file từ người dùng.")
    add_heading(doc, "7.6. Hạn chế", 2)
    add_bullets(doc, [
        "Hệ thống chưa tích hợp đăng nhập một lần SSO với hệ thống tài khoản của trường.",
        "Chức năng báo cáo thống kê còn ở mức cơ bản, chưa có phân tích xu hướng theo nhiều năm.",
        "Chưa có cơ chế ký số hoặc xác nhận điện tử cho quyết định khen thưởng cuối cùng.",
        "Việc quét file nén mới dừng ở mức ClamAV, chưa bổ sung cơ chế phân tích zip bomb chuyên sâu.",
    ])


def add_chapter_8(doc):
    add_heading(doc, "CHƯƠNG 8. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", 1)
    add_heading(doc, "8.1. Kết luận", 2)
    add_body(doc, "Khóa luận đã xây dựng thành công hệ thống xét duyệt thi đua khen thưởng cấp trường với các chức năng trọng tâm: quản lý người dùng, cấu hình tiêu chí, tạo và nộp hồ sơ, upload minh chứng an toàn, duyệt nhiều cấp, chấm điểm chi tiết, reassign reviewer, nhắc việc quá hạn, dashboard vận hành và audit log.")
    add_body(doc, "Hệ thống giúp số hóa quy trình thủ công, giảm sai sót trong xử lý hồ sơ, tăng tính minh bạch khi chấm điểm và hỗ trợ quản lý tập trung. Các kết quả kiểm thử cho thấy hệ thống đáp ứng được yêu cầu nghiệp vụ cốt lõi và có khả năng triển khai trong phạm vi thực tế.")
    add_heading(doc, "8.2. Đóng góp của đề tài", 2)
    add_bullets(doc, [
        "Đề xuất mô hình số hóa quy trình thi đua khen thưởng phù hợp môi trường đại học.",
        "Hiện thực hệ thống web đầy đủ từ frontend, backend đến database.",
        "Tích hợp bảo mật upload file bằng ClamAV và kiểm soát tải xuống.",
        "Bổ sung cơ chế chấm điểm theo ý nhỏ, giúp hội đồng đánh giá chi tiết và công bằng hơn.",
        "Cung cấp dashboard vận hành và audit log phục vụ quản trị.",
    ])
    add_heading(doc, "8.3. Hướng phát triển", 2)
    add_bullets(doc, [
        "Tích hợp SSO với tài khoản nhà trường để đồng bộ người dùng.",
        "Bổ sung báo cáo phân tích nâng cao theo khoa, năm học, loại hình thành tích và điểm trung bình.",
        "Phát triển chức năng xuất quyết định, giấy chứng nhận và biểu mẫu theo chuẩn hành chính.",
        "Tích hợp chữ ký số hoặc phê duyệt điện tử cho quyết định cuối cùng.",
        "Mở rộng cơ chế bảo mật file bằng CDR hoặc sandbox phân tích file nâng cao.",
        "Xây dựng ứng dụng mobile hoặc giao diện responsive tối ưu hơn cho người dùng thường xuyên.",
    ])
    add_heading(doc, "8.4. Đề xuất triển khai thực tế", 2)
    add_body(doc, "Để triển khai thực tế, hệ thống cần được đặt trên môi trường server ổn định, tách biệt cấu hình dev, staging và production. ClamAV nên chạy dưới dạng service tự khởi động cùng hệ thống. Database cần được backup định kỳ và kiểm tra khôi phục. Trước khi vận hành chính thức, cần tổ chức UAT với đại diện sinh viên, cán bộ, hội đồng và quản trị viên để chốt quy trình cuối cùng.")


def add_references_appendix(doc):
    doc.add_page_break()
    add_heading(doc, "TÀI LIỆU THAM KHẢO", 1)
    refs = [
        "React Documentation, https://react.dev",
        "Express.js Documentation, https://expressjs.com",
        "Prisma Documentation, https://www.prisma.io/docs",
        "PostgreSQL Documentation, https://www.postgresql.org/docs",
        "ClamAV Documentation, https://docs.clamav.net",
        "OWASP Foundation, OWASP File Upload Cheat Sheet, https://cheatsheetseries.owasp.org",
        "OWASP Foundation, OWASP Top 10 Web Application Security Risks.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"[{i}]. {ref}")
        set_run_font(run, size=13)
    doc.add_page_break()
    add_heading(doc, "PHỤ LỤC", 1)
    add_heading(doc, "Phụ lục A. Cấu hình môi trường chính", 2)
    add_table(doc, ["Biến môi trường", "Mục đích"], [
        ["DATABASE_URL", "Chuỗi kết nối PostgreSQL."],
        ["JWT_SECRET", "Khóa ký JWT."],
        ["MAIL_HOST, MAIL_PORT, MAIL_USER", "Cấu hình gửi email."],
        ["CLAMAV_HOST, CLAMAV_PORT", "Thông tin kết nối clamd."],
        ["ALLOW_UNSCANNED_UPLOADS", "Quyết định có cho upload khi không quét được hay không."],
    ], title="Bảng PL.1. Biến môi trường tiêu biểu")
    add_heading(doc, "Phụ lục B. Nhật ký kiểm thử UAT", 2)
    add_body(doc, "UAT đã kiểm tra thành công các luồng: đăng nhập theo vai trò, upload minh chứng sạch, tạo hồ sơ, nộp hồ sơ, duyệt cấp đơn vị, reassign cấp khoa, duyệt cấp khoa, chấm điểm cấp trường, kiểm tra trạng thái APPROVED, nhắc việc quá hạn và dashboard vận hành.")


def build():
    if OUT.exists():
        BACKUP.write_bytes(OUT.read_bytes())
    doc = Document()
    configure_document(doc)
    add_header_footer(doc)
    add_cover(doc)
    add_front_matter(doc)
    add_toc(doc)
    add_chapter_1(doc)
    doc.add_page_break()
    add_chapter_2(doc)
    doc.add_page_break()
    add_chapter_3(doc)
    doc.add_page_break()
    add_chapter_4(doc)
    doc.add_page_break()
    add_chapter_5(doc)
    doc.add_page_break()
    add_chapter_6(doc)
    doc.add_page_break()
    add_chapter_7(doc)
    doc.add_page_break()
    add_chapter_8(doc)
    add_references_appendix(doc)

    for paragraph in doc.paragraphs:
        set_paragraph_font(paragraph, size=13)
    doc.save(ACADEMIC_OUT)
    wrote = ACADEMIC_OUT
    try:
        doc.save(OUT)
        wrote = OUT
    except PermissionError:
        try:
            doc.save(ALT_OUT)
            wrote = ALT_OUT
        except PermissionError:
            try:
                doc.save(ALT_OUT_2)
                wrote = ALT_OUT_2
            except PermissionError:
                wrote = ACADEMIC_OUT
    print(f"WROTE {wrote}")
    print(f"ACADEMIC {ACADEMIC_OUT}")
    print(f"BACKUP {BACKUP}")


if __name__ == "__main__":
    build()
